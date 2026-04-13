"""
Generate KLAR User Manual in English and German as DOCX files.
Font: Arial Narrow throughout. Professional structure with proper formatting.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

FONT_NAME = "Arial Narrow"
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "docs")

def setup_styles(doc):
    """Configure all document styles to use Arial Narrow."""
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25
    rpr = style.element.rPr
    rpr_rfonts = rpr.find(qn("w:rFonts"))
    if rpr_rfonts is None:
        rpr_rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rpr_rfonts)
    rpr_rfonts.set(qn("w:eastAsia"), FONT_NAME)

    for level in range(1, 5):
        sname = f"Heading {level}"
        hs = doc.styles[sname]
        hs.font.name = FONT_NAME
        hs.font.color.rgb = RGBColor(0x00, 0x6E, 0x3E) if level <= 2 else RGBColor(0x22, 0x22, 0x22)
        hs.font.bold = True
        sizes = {1: 26, 2: 20, 3: 14, 4: 12}
        hs.font.size = Pt(sizes.get(level, 12))
        hs.paragraph_format.space_before = Pt(18 if level == 1 else 14)
        hs.paragraph_format.space_after = Pt(6)
        r = hs.element.rPr
        rf = r.find(qn("w:rFonts"))
        if rf is None:
            rf = r.makeelement(qn("w:rFonts"), {})
            r.insert(0, rf)
        rf.set(qn("w:eastAsia"), FONT_NAME)

    # Title style
    ts = doc.styles["Title"]
    ts.font.name = FONT_NAME
    ts.font.size = Pt(32)
    ts.font.bold = True
    ts.font.color.rgb = RGBColor(0x00, 0x8C, 0x50)
    ts.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts.paragraph_format.space_after = Pt(4)

    # Subtitle
    ss = doc.styles["Subtitle"]
    ss.font.name = FONT_NAME
    ss.font.size = Pt(14)
    ss.font.color.rgb = RGBColor(0x55, 0x55, 0x58)
    ss.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ss.paragraph_format.space_after = Pt(24)


def add_cover(doc, lang):
    """Add cover page."""
    for _ in range(6):
        doc.add_paragraph("")

    doc.add_paragraph("KLAR", style="Title")

    if lang == "en":
        doc.add_paragraph("Knowledge Legitimacy Audit & Review", style="Subtitle")
        p = doc.add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Complete User Manual")
        run.font.size = Pt(16)
        run.font.name = FONT_NAME
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p2 = doc.add_paragraph("")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("Version 1.0 — April 2026")
        r2.font.size = Pt(11)
        r2.font.name = FONT_NAME
        r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    else:
        doc.add_paragraph("Knowledge Legitimacy Audit & Review", style="Subtitle")
        p = doc.add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Vollständiges Benutzerhandbuch")
        run.font.size = Pt(16)
        run.font.name = FONT_NAME
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p2 = doc.add_paragraph("")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("Version 1.0 — April 2026")
        r2.font.size = Pt(11)
        r2.font.name = FONT_NAME
        r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = FONT_NAME
        r.font.size = Pt(11)
        p.add_run(text).font.name = FONT_NAME
    else:
        r = p.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = Pt(11)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = FONT_NAME
        r.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.name = FONT_NAME
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = Pt(11)


def add_colored_bullet(doc, color_hex, text):
    """Add a bullet with a colored prefix."""
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("● ")
    r.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    )
    r.font.name = FONT_NAME
    r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.name = FONT_NAME
    r2.font.size = Pt(11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = FONT_NAME
                r.font.size = Pt(10)
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT_NAME
                    r.font.size = Pt(10)
    doc.add_paragraph("")


# ─────────────────────────────────────────────────────────────────────────
# ENGLISH MANUAL
# ─────────────────────────────────────────────────────────────────────────
def build_english(doc):
    setup_styles(doc)
    add_cover(doc, "en")

    # TABLE OF CONTENTS placeholder
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. The Problem: Why KLAR Exists",
        "2. The Solution: What KLAR Does",
        "3. Features at a Glance",
        "4. Detailed Feature Guide",
        "   4.1 Fact Checking",
        "   4.2 AI Detection",
        "   4.3 Bias Detection",
        "   4.4 Plagiarism Check",
        "   4.5 Quality Evaluation",
        "   4.6 File Upload",
        "   4.7 URL Verification",
        "   4.8 Trust Report",
        "   4.9 Dashboard & History",
        "   4.10 Human Review System",
        "5. How the Technology Works",
        "6. The Web Application (UI Guide)",
        "7. The Chrome Extension",
        "8. Browser Bookmarklets",
        "9. API Access for Developers",
        "10. Pricing & Usage Limits",
        "11. Privacy, Security & GDPR",
        "12. Who Is KLAR For?",
        "13. Future Roadmap",
        "14. Getting Started (Quick Start)",
        "15. Frequently Asked Questions",
    ]
    for item in toc_items:
        add_para(doc, item)
    doc.add_page_break()

    # ── 1. THE PROBLEM ──
    doc.add_heading("1. The Problem: Why KLAR Exists", level=1)
    add_para(doc, "We live in an age where artificial intelligence can generate thousands of pages of text in seconds. Tools like ChatGPT, Google Gemini, Claude, and others produce content that reads perfectly — grammatically correct, well-structured, and convincing.")
    add_para(doc, "But there is a fundamental problem:")
    doc.add_heading("AI makes things up.", level=3)
    add_para(doc, "This is called \"hallucination.\" AI models do not actually know facts. They predict what words are likely to follow other words based on patterns learned from training data. This means they regularly:")
    add_bullet(doc, "Invent statistics that do not exist")
    add_bullet(doc, "Cite sources that were never published")
    add_bullet(doc, "State historical dates or figures incorrectly")
    add_bullet(doc, "Present opinions as established facts")
    add_bullet(doc, "Mix true and false information in the same paragraph")
    add_para(doc, "Studies estimate that 40–73% of AI-generated texts contain at least one factual error. The danger is that these errors are invisible — the text sounds confident and authoritative, even when it is wrong.")
    doc.add_heading("Who is affected?", level=3)
    add_bullet(doc, "Journalists who receive AI-written press releases filled with fabricated claims")
    add_bullet(doc, "Students who rely on AI-generated study materials without knowing which facts are real")
    add_bullet(doc, "Businesses that use AI to draft reports, proposals, or documentation for clients")
    add_bullet(doc, "Researchers who need to validate data and detect plagiarism")
    add_bullet(doc, "Anyone who reads content online and wants to know if it is actually true")
    add_para(doc, "Manual fact-checking is slow, tedious, and impractical for the volume of AI-generated content being produced today. You would need to search each claim individually, cross-reference multiple sources, and document your findings — a process that could take hours for a single article.")
    add_para(doc, "That is exactly the problem KLAR solves.")
    doc.add_page_break()

    # ── 2. THE SOLUTION ──
    doc.add_heading("2. The Solution: What KLAR Does", level=1)
    add_para(doc, "KLAR stands for Knowledge Legitimacy Audit & Review. It is a web application that takes any text — whether AI-generated or human-written — and automatically:")
    add_bullet(doc, "Finds every factual claim in the text", bold_prefix="Extracts claims: ")
    add_bullet(doc, "Searches Wikipedia, Wikidata, and Google for evidence", bold_prefix="Searches for evidence: ")
    add_bullet(doc, "Classifies each claim as Supported, Contradicted, or Unconfirmed", bold_prefix="Judges each claim: ")
    add_bullet(doc, "Provides source links for every verdict", bold_prefix="Cites sources: ")
    add_bullet(doc, "Determines if the text was written by AI", bold_prefix="Detects AI origin: ")
    add_bullet(doc, "Identifies emotional language, one-sided framing, and political lean", bold_prefix="Analyzes bias: ")
    add_bullet(doc, "Checks for copied content from the web", bold_prefix="Finds plagiarism: ")
    add_bullet(doc, "Scores the text using professional evaluation frameworks", bold_prefix="Evaluates quality: ")
    add_para(doc, "The result is a comprehensive Trust Report with a color-coded visualization of your text, where each claim is highlighted in green (supported), red (contradicted), or gray (unconfirmed). You can click on any claim to see the evidence and reasoning behind the verdict.")
    add_para(doc, "Think of KLAR as a smart research assistant that reads your text, questions every factual statement, finds the evidence for or against it, and presents you with a clear, transparent report — all in under 15 seconds.")
    doc.add_page_break()

    # ── 3. FEATURES AT A GLANCE ──
    doc.add_heading("3. Features at a Glance", level=1)
    features_table = [
        ["Fact Checking", "Verify every claim against Wikipedia & Google", "All plans"],
        ["AI Detection", "Detect if text was written by AI (0–100%)", "Free+"],
        ["Bias Detection", "Find emotional language & one-sided framing", "Free+"],
        ["Plagiarism Check", "Compare text against web sources", "Pro+"],
        ["Quality Evaluation", "MECE, Red Team, BLUF, Pre-Mortem scoring", "Pro+"],
        ["File Upload", "PDF, DOCX, TXT — up to 10 MB", "Free+"],
        ["URL Verification", "Analyze any web article by URL", "Free+"],
        ["Trust Report", "Color-coded report with source links", "All plans"],
        ["PDF Export", "Download the full report as PDF", "Pro+"],
        ["Dashboard", "Verification stats, trends, recent reports", "Free+"],
        ["History", "Searchable, sortable list of all verifications", "Free+"],
        ["Human Review", "Disagree with a verdict? Submit your review", "Free+"],
        ["Chrome Extension", "Right-click to verify text on any webpage", "Pro+"],
        ["Bookmarklets", "One-click verification from any browser", "All plans"],
        ["API Access", "REST API with SSE streaming", "Team"],
        ["Dark Mode", "Full dark/light/system theme support", "All plans"],
        ["Multilingual", "German & English interface", "All plans"],
        ["GDPR Controls", "Export data, delete account", "All plans"],
    ]
    add_table(doc, ["Feature", "Description", "Availability"], features_table)
    doc.add_page_break()

    # ── 4. DETAILED FEATURE GUIDE ──
    doc.add_heading("4. Detailed Feature Guide", level=1)

    doc.add_heading("4.1 Fact Checking — The Core Engine", level=2)
    add_para(doc, "Fact checking is the heart of KLAR. When you submit a text, here is exactly what happens:")
    add_bullet(doc, "KLAR uses Google Gemini Flash AI to read your text and identify every sentence that contains a factual claim — a statement that can be verified as true or false.", bold_prefix="Step 1 — Claim Extraction: ")
    add_bullet(doc, "For each claim, KLAR searches Wikipedia and Wikidata first (these cover about 40% of typical claims). For remaining claims, it uses Serper.dev to search the broader web. Each claim gets 3–5 sources.", bold_prefix="Step 2 — Evidence Search: ")
    add_bullet(doc, "A second AI pass evaluates each claim against the collected evidence and classifies it:", bold_prefix="Step 3 — Judgment: ")
    add_colored_bullet(doc, "16A34A", "Supported (Green) — The evidence confirms this claim is accurate")
    add_colored_bullet(doc, "DC2626", "Contradicted (Red) — The evidence shows this claim is wrong")
    add_colored_bullet(doc, "64748B", "Unconfirmed (Gray) — No reliable source could be found to confirm or deny this. This does not mean the claim is wrong.")
    add_para(doc, "Each verdict comes with a plain-English explanation of why the claim was judged that way, plus clickable links to every source used.")

    doc.add_heading("4.2 AI Detection", level=2)
    add_para(doc, "This engine analyzes writing patterns to determine whether text was written by a human or by an AI. It does not use AI to detect AI — instead, it uses statistical mathematics (NLP), which makes it instant and free to run.")
    add_para(doc, "What it measures:")
    add_bullet(doc, "AI tends to write very uniform sentences", bold_prefix="Sentence variety — ")
    add_bullet(doc, "AI uses a narrower range of words", bold_prefix="Vocabulary diversity — ")
    add_bullet(doc, "AI writing follows more predictable patterns", bold_prefix="Predictability — ")
    add_bullet(doc, "AI often uses specific phrases and structures", bold_prefix="Structural patterns — ")
    add_para(doc, "The result is a percentage score (0–100%) and a verdict: Human written, Likely human, Mixed, Likely AI, or AI generated.")

    doc.add_heading("4.3 Bias Detection", level=2)
    add_para(doc, "The bias detector reads your entire text and looks for signs of partiality — when writing tries to push the reader toward a particular opinion rather than presenting facts neutrally.")
    add_para(doc, "What it detects:")
    add_bullet(doc, "Words designed to trigger emotions (\"devastating\", \"miraculous\", \"shocking\")", bold_prefix="Loaded language — ")
    add_bullet(doc, "Trying to make the reader feel rather than think", bold_prefix="Emotional appeals — ")
    add_bullet(doc, "Presenting only one perspective on a topic", bold_prefix="One-sided framing — ")
    add_bullet(doc, "Relying too heavily on sources from one viewpoint", bold_prefix="Source imbalance — ")
    add_bullet(doc, "Whether the text leans left, right, or stays neutral", bold_prefix="Political lean — ")
    add_para(doc, "The result is a bias score from 0 to 100 (lower is better), a level classification (Minimal, Low, Moderate, Significant, Extreme), and a summary explaining what biases were found.")

    doc.add_heading("4.4 Plagiarism Check", level=2)
    add_para(doc, "The plagiarism checker compares your text against known web sources to identify copied or closely paraphrased content.")
    add_para(doc, "The result includes:")
    add_bullet(doc, "An originality percentage (higher is better, e.g., \"92% original\")")
    add_bullet(doc, "A list of any matching sources found online")
    add_bullet(doc, "A verdict: Original, Mostly Original, Some Overlap, Significant Overlap, or Likely Plagiarized")

    doc.add_heading("4.5 Quality Evaluation", level=2)
    add_para(doc, "This engine grades the quality of your text using four professional analysis frameworks that experts use to evaluate reports and documents:")
    add_bullet(doc, "Is the content well-organized with no gaps or overlaps?", bold_prefix="MECE (Mutually Exclusive, Collectively Exhaustive) — ")
    add_bullet(doc, "Could someone easily poke holes in the arguments?", bold_prefix="Red Team — ")
    add_bullet(doc, "Is the main point stated clearly at the start?", bold_prefix="BLUF (Bottom Line Up Front) — ")
    add_bullet(doc, "What could go wrong if someone acted on this information?", bold_prefix="Pre-Mortem — ")
    add_para(doc, "The result is an overall letter grade (A through F), individual scores for each framework, and specific recommendations for improvement.")

    doc.add_heading("4.6 File Upload", level=2)
    add_para(doc, "Instead of pasting text, you can upload a document directly:")
    add_bullet(doc, "PDF files")
    add_bullet(doc, "Word documents (DOCX)")
    add_bullet(doc, "Text files (TXT)")
    add_para(doc, "KLAR automatically extracts the text from your file and runs all the analyses you select. Maximum file size: 10 MB (2 MB on the Free plan).")

    doc.add_heading("4.7 URL Verification", level=2)
    add_para(doc, "Instead of copy-pasting an article, enter the web address (URL) of any article or blog post. KLAR visits the page, extracts the main content, and analyzes it. Supports news articles, blog posts, documentation pages, and most publicly accessible web pages.")

    doc.add_heading("4.8 Trust Report", level=2)
    add_para(doc, "After analysis, KLAR generates a comprehensive trust report that includes:")
    add_bullet(doc, "A trust score (0–100) based on the ratio of supported vs. contradicted claims (unconfirmed claims are excluded from the score)")
    add_bullet(doc, "Color-coded claims with verdicts and sources")
    add_bullet(doc, "A highlighted text view — your original text with claims marked in green (supported), red (contradicted), and gray (unconfirmed)")
    add_bullet(doc, "Source links for every verdict")
    add_bullet(doc, "PDF export — download the full report")

    doc.add_heading("4.9 Dashboard & Verification History", level=2)
    add_para(doc, "The Dashboard gives you a personal overview:")
    add_bullet(doc, "How many verifications you have run")
    add_bullet(doc, "Your trust score trends over time")
    add_bullet(doc, "Recent contradicted claims to watch out for")
    add_bullet(doc, "Quick access to your latest reports")
    add_para(doc, "The Verification History is a searchable, sortable list of every text you have ever verified. Filter by date, sort by trust score, search by keywords. Click any entry to re-open the full report.\n\nNote: The trust score measures supported claims vs. contradicted claims. Unconfirmed claims (where no source could be found either way) are excluded from the score — they do not count against you.")

    doc.add_heading("4.10 Human Review System", level=2)
    add_para(doc, "If you disagree with KLAR's verdict on a claim, you can submit a review with your own assessment and comments. This creates a record that helps improve future accuracy.")
    doc.add_page_break()

    # ── 5. HOW THE TECHNOLOGY WORKS ──
    doc.add_heading("5. How the Technology Works", level=1)
    add_para(doc, "KLAR does not use a single AI that does everything. Instead, it uses a pipeline — a series of specialized steps that each do one thing well, working together like an assembly line.")
    doc.add_heading("The 10-Step Pipeline", level=2)
    pipeline_steps = [
        ("1. AI Detection", "Statistical analysis determines if the text is human or AI-written (instant, no AI needed)."),
        ("2. Quality Evaluation", "Text is scored against MECE, Red Team, BLUF, and Pre-Mortem frameworks."),
        ("3. Claim Extraction", "Gemini Flash AI reads the text and identifies every factual claim as structured data."),
        ("4. Evidence Search", "Each claim is searched on Wikipedia/Wikidata first, then Google for remaining claims. 3–5 sources per claim."),
        ("5. Claim Quality Check", "Each claim is rated for how specific and verifiable it is."),
        ("6. Hallucination Detection", "Statistical patterns typical of AI fabrication are identified."),
        ("7. AI Judgment", "A second Gemini pass evaluates each claim against the collected evidence."),
        ("8. Bias Analysis", "NLP-based detection of emotional language, framing, and source balance."),
        ("9. Plagiarism Check", "Text fingerprints are compared against web search results."),
        ("10. Cross-Reference", "Claims are validated across multiple sources for consistency."),
    ]
    for step_title, step_desc in pipeline_steps:
        add_bullet(doc, step_desc, bold_prefix=step_title + " — ")

    doc.add_heading("Which AI Does KLAR Use?", level=2)
    add_para(doc, "KLAR uses Google Gemini 2.5 Flash — one of Google's latest AI models. It uses it in two ways:")
    add_bullet(doc, "For extracting claims and understanding text", bold_prefix="Standard mode — ")
    add_bullet(doc, "Gemini searches the live internet in real-time (\"Google Search Grounding\")", bold_prefix="Grounded Search mode — ")
    add_para(doc, "KLAR does NOT use vector databases, agent frameworks, RAG (retrieval-augmented generation), or fine-tuned models. It searches live every time for the most current information.")

    doc.add_heading("The Analysis Engines", level=2)
    engines_table = [
        ["Fact Checker", "AI + Search", "Gemini extracts claims → source search → Gemini judges"],
        ["AI Detector", "Statistical Math", "Sentence patterns, word variety, predictability"],
        ["Bias Detector", "Statistical Math", "Emotional words, framing patterns, source balance"],
        ["Plagiarism Checker", "Statistical Math", "Text fingerprints vs web search results"],
        ["Quality Evaluator", "Statistical Math", "MECE/Red Team/BLUF/Pre-Mortem scoring"],
        ["Hallucination Detector", "Statistical Math", "AI fabrication patterns detection"],
        ["Cross-Reference", "Logic", "Consistency checks across multiple sources"],
    ]
    add_table(doc, ["Engine", "Type", "How It Works"], engines_table)

    doc.add_heading("Front-End vs. Back-End", level=2)
    add_para(doc, "KLAR has two parts that work separately:")
    add_bullet(doc, "Built with Next.js and React. Runs in your browser. Handles the interface, animations, dark mode, and language switching. Does NOT run any AI.", bold_prefix="Front-End (what you see): ")
    add_bullet(doc, "Runs on the server. Calls Google Gemini AI, searches Wikipedia and Google, runs all analysis engines, and stores your reports in a Supabase database (Frankfurt, EU).", bold_prefix="Back-End (behind the scenes): ")
    add_para(doc, "When you click \"Analyze,\" the front-end sends your text to the back-end API. The back-end processes it step by step and streams results back in real-time — so you see each claim being verified live as it happens.")
    doc.add_page_break()

    # ── 6. THE WEB APPLICATION (UI Guide) ──
    doc.add_heading("6. The Web Application (UI Guide)", level=1)
    add_para(doc, "KLAR is accessible at klar-app.vercel.app. The interface supports both German and English, with a language toggle in the header. Dark mode, light mode, and system-preference detection are built in.")

    doc.add_heading("Main Pages", level=2)
    add_bullet(doc, "Modern landing page explaining what KLAR does and how to get started.", bold_prefix="Landing Page — ")
    add_bullet(doc, "The main analysis interface. Paste text, upload a file, or enter a URL. Select which analysis modes to run. Click \"Analyze.\"", bold_prefix="Verify Page — ")
    add_bullet(doc, "Personal overview with verification stats, trust score trends, and recent reports.", bold_prefix="Dashboard — ")
    add_bullet(doc, "Searchable, sortable list of all past verifications. Filter by date, score, or keywords.", bold_prefix="History — ")
    add_bullet(doc, "Profile management, data export (GDPR), and account deletion.", bold_prefix="Settings — ")

    doc.add_heading("The Verification Flow", level=2)
    add_bullet(doc, "Navigate to the Verify page")
    add_bullet(doc, "Paste your text (up to 10,000 characters), upload a file, or enter a URL")
    add_bullet(doc, "Select analysis modes (Fact Check is default; add AI Detection, Bias, Plagiarism, Quality, or Comprehensive)")
    add_bullet(doc, "Click \"Analyze\"")
    add_bullet(doc, "Watch results stream in real-time — each claim appears as it is verified")
    add_bullet(doc, "Review the Trust Report with color-coded verdicts")
    add_bullet(doc, "Click any claim to expand its evidence, sources, and reasoning")
    add_bullet(doc, "Export as PDF if needed")

    doc.add_heading("Authentication", level=2)
    add_para(doc, "KLAR supports two authentication methods:")
    add_bullet(doc, "Standard signup with email and password", bold_prefix="Email + Password — ")
    add_bullet(doc, "One-click sign-in with your Google account", bold_prefix="Google OAuth — ")
    add_para(doc, "Guests can use KLAR without signing up (3 checks/month). Creating a free account gives you 10 checks/month and access to more features.")
    doc.add_page_break()

    # ── 7. CHROME EXTENSION ──
    doc.add_heading("7. The Chrome Extension", level=1)
    add_para(doc, "KLAR includes a Manifest V3 Chrome Extension that lets you verify text directly on any webpage without leaving the site.")
    doc.add_heading("How to Install", level=2)
    add_bullet(doc, "Get an API key from your KLAR Settings page (klar-app.vercel.app/en/settings)")
    add_bullet(doc, "Open chrome://extensions in your browser")
    add_bullet(doc, "Enable \"Developer mode\" (toggle in top-right corner)")
    add_bullet(doc, "Click \"Load unpacked\" and select the extension/ folder from the KLAR repository")
    add_bullet(doc, "Click the KLAR extension icon in your toolbar and enter your API key")
    doc.add_heading("How to Use", level=2)
    add_bullet(doc, "Select text on any webpage, right-click, and choose \"Verify with KLAR\"", bold_prefix="Right-click → Verify: ")
    add_bullet(doc, "Right-click anywhere on a page (without selecting text) to verify the entire page", bold_prefix="Whole page verification: ")
    add_bullet(doc, "Results appear in a floating panel directly on the page, showing trust score and claim verdicts. Only sources the AI actually referenced are shown.", bold_prefix="Inline results: ")
    add_bullet(doc, "The extension automatically adapts to your system theme", bold_prefix="Dark mode: ")
    add_para(doc, "The extension is also compatible with Edge and all Chromium-based browsers.")
    doc.add_page_break()

    # ── 8. BOOKMARKLETS ──
    doc.add_heading("8. Browser Bookmarklets", level=1)
    add_para(doc, "For browsers that do not support Chrome extensions (or if you prefer a simpler approach), KLAR provides bookmarklets — small buttons you drag to your browser's bookmarks bar:")
    add_bullet(doc, "Sends the selected text or entire page content to KLAR for analysis", bold_prefix="\"KLAR Verify\" — ")
    add_bullet(doc, "Sends the current page URL to KLAR for URL-based verification", bold_prefix="\"KLAR URL\" — ")
    add_para(doc, "No installation needed. Works in Chrome, Firefox, Edge, and Safari.")
    doc.add_page_break()

    # ── 9. API ACCESS ──
    doc.add_heading("9. API Access for Developers", level=1)
    add_para(doc, "Developers can integrate KLAR into their own applications using a REST API. Available on the Team plan.")
    add_para(doc, "How it works:")
    add_bullet(doc, "Send a POST request to /api/verify with your text and API key")
    add_bullet(doc, "Results are streamed back in real-time using Server-Sent Events (SSE)")
    add_bullet(doc, "The Chrome Extension uses a separate /api/extension/scan endpoint that returns JSON")
    add_para(doc, "Code examples are available in cURL, JavaScript, and Python.")
    doc.add_page_break()

    # ── 10. PRICING ──
    doc.add_heading("10. Pricing & Usage Limits", level=1)
    pricing_table = [
        ["Guest", "Free (no signup)", "3/month", "2,000 chars", "No", "Fact Check only"],
        ["Free", "$0/month", "10/month", "5,000 chars", "Up to 2 MB", "Fact Check, Bias, AI"],
        ["Pro", "$9/month", "200/month", "10,000 chars", "Up to 10 MB", "All 5 + Comprehensive"],
        ["Team", "$29/user/month", "Unlimited", "10,000 chars", "Up to 10 MB", "All + API + Concurrent"],
    ]
    add_table(doc, ["Plan", "Price", "Monthly Checks", "Max Text", "File Upload", "Analysis Modes"], pricing_table)
    doc.add_heading("Usage Notes", level=2)
    add_bullet(doc, "Comprehensive analysis (all 5 engines at once) counts as 2–3 checks depending on your plan")
    add_bullet(doc, "Free and guest users can run one analysis at a time")
    add_bullet(doc, "Burst protection: maximum 1–2 checks per minute on free plans")
    add_bullet(doc, "Abuse protection: repeatedly hitting limits triggers escalating cooldowns (5 min → 30 min → 2 hours)")
    doc.add_page_break()

    # ── 11. PRIVACY & SECURITY ──
    doc.add_heading("11. Privacy, Security & GDPR", level=1)
    add_para(doc, "KLAR is built with European data protection standards at its core:")
    add_bullet(doc, "Built for the EU General Data Protection Regulation", bold_prefix="GDPR-compliant — ")
    add_bullet(doc, "All data stored in Frankfurt, Germany (Supabase EU)", bold_prefix="EU data hosting — ")
    add_bullet(doc, "Transparent AI decision-making", bold_prefix="EU AI Act ready — ")
    add_bullet(doc, "No tracking cookies, no data selling, no third-party analytics", bold_prefix="No surveillance — ")
    add_bullet(doc, "Each user can only see their own data (Row Level Security)", bold_prefix="Data isolation — ")
    add_bullet(doc, "Download everything KLAR has stored about you at any time", bold_prefix="Data export — ")
    add_bullet(doc, "Permanently remove your account and all associated data", bold_prefix="Account deletion — ")
    doc.add_heading("Security Measures", level=2)
    add_bullet(doc, "Row Level Security (RLS) on all Supabase tables")
    add_bullet(doc, "API keys stored server-side only — never exposed to browsers")
    add_bullet(doc, "Input validation and sanitization on all text inputs")
    add_bullet(doc, "Prompt injection prevention with isolated AI calls")
    add_bullet(doc, "CSRF protection and HTTPS enforcement")
    add_bullet(doc, "Rate limiting per user and IP address")
    add_bullet(doc, "Content Security Policy headers")
    doc.add_page_break()

    # ── 12. WHO IS KLAR FOR? ──
    doc.add_heading("12. Who Is KLAR For?", level=1)
    add_bullet(doc, "Verify claims in press releases, statements, and articles before publishing", bold_prefix="Journalists — ")
    add_bullet(doc, "Check AI-generated study materials and research papers for accuracy", bold_prefix="Students — ")
    add_bullet(doc, "Validate data, detect plagiarism, and ensure academic integrity", bold_prefix="Researchers — ")
    add_bullet(doc, "Ensure blog posts, social media content, and articles are factually accurate", bold_prefix="Content Creators — ")
    add_bullet(doc, "Audit AI-generated reports, proposals, and client documentation", bold_prefix="Businesses — ")
    add_bullet(doc, "Verify contracts, compliance documents, and legal research", bold_prefix="Legal & Consulting Firms — ")
    add_bullet(doc, "Ensure teaching materials and student submissions are accurate and original", bold_prefix="Educational Institutions — ")
    add_bullet(doc, "Anyone who wants to know if what they are reading is actually true", bold_prefix="Everyone — ")
    doc.add_page_break()

    # ── 13. FUTURE ROADMAP ──
    doc.add_heading("13. Future Roadmap", level=1)
    add_para(doc, "KLAR is actively developed. Here is what we are planning for future versions:")
    add_bullet(doc, "Real-time monitoring of news sources and social media for spreading misinformation", bold_prefix="Misinformation monitoring — ")
    add_bullet(doc, "Support for larger teams with admin dashboards and usage analytics", bold_prefix="Team collaboration — ")
    add_bullet(doc, "Build custom verification workflows using KLAR's API", bold_prefix="Advanced API features — ")
    add_bullet(doc, "Extensions for Firefox, Safari, and other browsers", bold_prefix="More browser extensions — ")
    add_bullet(doc, "Analyze images and videos for manipulated content", bold_prefix="Visual content verification — ")
    add_bullet(doc, "Even more interface languages beyond German and English", bold_prefix="Additional languages — ")
    add_bullet(doc, "Integration with Slack, Teams, and other communication tools", bold_prefix="Workspace integrations — ")
    doc.add_page_break()

    # ── 14. GETTING STARTED ──
    doc.add_heading("14. Getting Started (Quick Start)", level=1)
    doc.add_heading("Option A: Instant Check (No Signup)", level=2)
    add_bullet(doc, "Go to klar-app.vercel.app")
    add_bullet(doc, "Click \"Verify\" in the navigation")
    add_bullet(doc, "Paste any text into the text box")
    add_bullet(doc, "Click \"Analyze\"")
    add_bullet(doc, "Review your Trust Report")
    add_para(doc, "You get 3 free checks per month as a guest.")

    doc.add_heading("Option B: Create a Free Account", level=2)
    add_bullet(doc, "Click \"Sign Up\" and enter your email + password (or use Google)")
    add_bullet(doc, "Verify your email address")
    add_bullet(doc, "You now get 10 checks/month, a personal dashboard, and verification history")

    doc.add_heading("Option C: Install the Chrome Extension", level=2)
    add_bullet(doc, "Get an API key from Settings (requires a free account)")
    add_bullet(doc, "Load the extension in Chrome's developer mode")
    add_bullet(doc, "Right-click any text on any webpage to verify it instantly")
    doc.add_page_break()

    # ── 15. FAQ ──
    doc.add_heading("15. Frequently Asked Questions", level=1)
    faqs = [
        ("How accurate is KLAR?", "KLAR's accuracy depends on the availability of sources. Claims that can be verified against Wikipedia and reliable web sources receive highly accurate verdicts. Claims about very recent events or niche topics may be classified as 'Unconfirmed' when no reliable source exists. Unconfirmed does not mean wrong — it simply means no source could be found to confirm or deny it."),
        ("Does KLAR store my original text?", "KLAR stores the extracted claims and their verdicts so you can access your history. All data is stored in Frankfurt (EU) with Row Level Security."),
        ("Can KLAR check text in languages other than German and English?", "The analysis engine processes text in many languages since the source search works across languages. The user interface is available in German and English."),
        ("Is my data safe?", "Yes. KLAR is GDPR-compliant, stores all data in Frankfurt (EU), uses Row Level Security for data isolation, and does not use tracking cookies or sell data to third parties."),
        ("How is KLAR different from just asking ChatGPT to fact-check?", "ChatGPT cannot access the internet in real-time and often makes up sources. KLAR actually searches live sources (Wikipedia, Google) and provides clickable links to every source it uses. Every claim is verified against real, external evidence."),
        ("Can I use KLAR commercially?", "Yes. The Pro and Team plans are designed for professional and commercial use."),
    ]
    for q, a in faqs:
        doc.add_heading(q, level=3)
        add_para(doc, a)

    add_para(doc, "")
    p = doc.add_paragraph("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Manual —")
    r.font.name = FONT_NAME
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r.italic = True


# ─────────────────────────────────────────────────────────────────────────
# GERMAN MANUAL
# ─────────────────────────────────────────────────────────────────────────
def build_german(doc):
    setup_styles(doc)
    add_cover(doc, "de")

    doc.add_heading("Inhaltsverzeichnis", level=1)
    toc_items = [
        "1. Das Problem: Warum es KLAR gibt",
        "2. Die Lösung: Was KLAR macht",
        "3. Features im Überblick",
        "4. Ausführliche Feature-Beschreibung",
        "   4.1 Faktencheck",
        "   4.2 AI-Erkennung",
        "   4.3 Bias-Erkennung",
        "   4.4 Plagiatscheck",
        "   4.5 Qualitätsbewertung",
        "   4.6 Datei-Upload",
        "   4.7 URL-Verifizierung",
        "   4.8 Trust Report",
        "   4.9 Dashboard & Historie",
        "   4.10 Menschliches Bewertungssystem",
        "5. Wie die Technologie funktioniert",
        "6. Die Web-Anwendung (UI-Leitfaden)",
        "7. Die Chrome Extension",
        "8. Browser-Bookmarklets",
        "9. API-Zugang für Entwickler",
        "10. Preise & Nutzungslimits",
        "11. Datenschutz, Sicherheit & DSGVO",
        "12. Für wen ist KLAR?",
        "13. Zukünftige Entwicklung",
        "14. Schnellstart-Anleitung",
        "15. Häufige Fragen",
    ]
    for item in toc_items:
        add_para(doc, item)
    doc.add_page_break()

    # ── 1 ──
    doc.add_heading("1. Das Problem: Warum es KLAR gibt", level=1)
    add_para(doc, "Wir leben in einer Zeit, in der künstliche Intelligenz in Sekunden Tausende von Textseiten erzeugen kann. Werkzeuge wie ChatGPT, Google Gemini, Claude und andere produzieren Inhalte, die perfekt klingen — grammatisch korrekt, gut strukturiert und überzeugend.")
    add_para(doc, "Aber es gibt ein grundlegendes Problem:")
    doc.add_heading("KI erfindet Dinge.", level=3)
    add_para(doc, 'Das nennt man \u201eHalluzination\u201c. KI-Modelle kennen keine Fakten. Sie sagen voraus, welche W\u00f6rter wahrscheinlich auf andere W\u00f6rter folgen, basierend auf Mustern aus Trainingsdaten. Das bedeutet, sie:')
    add_bullet(doc, "Erfinden Statistiken, die nicht existieren")
    add_bullet(doc, "Zitieren Quellen, die nie veröffentlicht wurden")
    add_bullet(doc, "Geben historische Daten oder Zahlen falsch an")
    add_bullet(doc, "Präsentieren Meinungen als etablierte Fakten")
    add_bullet(doc, "Mischen wahre und falsche Informationen im selben Absatz")
    add_para(doc, "Studien zeigen, dass 40–73 % aller KI-generierten Texte mindestens einen Faktenfehler enthalten. Die Gefahr ist, dass diese Fehler unsichtbar sind — der Text klingt selbstbewusst und autoritativ, auch wenn er falsch ist.")
    doc.add_heading("Wer ist betroffen?", level=3)
    add_bullet(doc, "Journalisten, die KI-geschriebene Pressemitteilungen mit erfundenen Behauptungen erhalten")
    add_bullet(doc, "Studierende, die sich auf KI-generierte Lernmaterialien verlassen")
    add_bullet(doc, "Unternehmen, die KI für Berichte, Angebote oder Dokumentationen nutzen")
    add_bullet(doc, "Forscher, die Daten validieren und Plagiate erkennen müssen")
    add_bullet(doc, "Jeder, der online Inhalte liest und wissen will, ob sie wahr sind")
    add_para(doc, "Manuelles Faktenprüfen ist langsam, mühsam und bei der Menge an KI-generierten Inhalten unpraktisch. Man müsste jede Behauptung einzeln suchen, mehrere Quellen abgleichen und die Ergebnisse dokumentieren — ein Prozess, der für einen einzelnen Artikel Stunden dauern kann.")
    add_para(doc, "Genau dieses Problem löst KLAR.")
    doc.add_page_break()

    # ── 2 ──
    doc.add_heading("2. Die Lösung: Was KLAR macht", level=1)
    add_para(doc, "KLAR steht für Knowledge Legitimacy Audit & Review. Es ist eine Web-Anwendung, die jeden Text — ob KI-generiert oder menschlich geschrieben — automatisch:")
    add_bullet(doc, "Findet jede faktische Behauptung im Text", bold_prefix="Behauptungen extrahieren: ")
    add_bullet(doc, "Durchsucht Wikipedia, Wikidata und Google nach Belegen", bold_prefix="Belege suchen: ")
    add_bullet(doc, "Stuft jede Behauptung ein als Bestätigt, Widerlegt oder Unbestätigt", bold_prefix="Jede Behauptung bewerten: ")
    add_bullet(doc, "Liefert Quellenlinks für jedes Urteil", bold_prefix="Quellen zitieren: ")
    add_bullet(doc, "Erkennt, ob der Text von KI geschrieben wurde", bold_prefix="KI-Herkunft erkennen: ")
    add_bullet(doc, "Identifiziert emotionale Sprache, einseitige Darstellung und politische Tendenz", bold_prefix="Bias analysieren: ")
    add_bullet(doc, "Prüft auf kopierten Inhalt aus dem Web", bold_prefix="Plagiate finden: ")
    add_bullet(doc, "Bewertet den Text mit professionellen Analyse-Frameworks", bold_prefix="Qualität bewerten: ")
    add_para(doc, "Das Ergebnis ist ein umfassender Trust Report mit einer farbcodierten Darstellung Ihres Textes, in der jede Behauptung in Grün (bestätigt), Rot (widerlegt) oder Grau (unbestätigt) markiert ist. Sie können auf jede Behauptung klicken, um die Belege und die Begründung hinter dem Urteil zu sehen.")
    add_para(doc, "Stellen Sie sich KLAR als einen intelligenten Forschungsassistenten vor, der Ihren Text liest, jede faktische Aussage hinterfragt, die Belege dafür oder dagegen findet und Ihnen einen klaren, transparenten Bericht präsentiert — alles in unter 15 Sekunden.")
    doc.add_page_break()

    # ── 3 ──
    doc.add_heading("3. Features im Überblick", level=1)
    features_table = [
        ["Faktencheck", "Jede Behauptung gegen Wikipedia & Google prüfen", "Alle Pläne"],
        ["AI-Erkennung", "Erkennen, ob Text von KI geschrieben wurde (0–100 %)", "Free+"],
        ["Bias-Erkennung", "Emotionale Sprache & einseitige Darstellung finden", "Free+"],
        ["Plagiatscheck", "Text gegen Webquellen vergleichen", "Pro+"],
        ["Qualitätsbewertung", "MECE, Red Team, BLUF, Pre-Mortem Bewertung", "Pro+"],
        ["Datei-Upload", "PDF, DOCX, TXT — bis zu 10 MB", "Free+"],
        ["URL-Verifizierung", "Jeden Web-Artikel per URL analysieren", "Free+"],
        ["Trust Report", "Farbcodierter Bericht mit Quellenlinks", "Alle Pläne"],
        ["PDF-Export", "Vollständigen Bericht als PDF herunterladen", "Pro+"],
        ["Dashboard", "Statistiken, Trends, aktuelle Berichte", "Free+"],
        ["Historie", "Durchsuchbare, sortierbare Liste aller Prüfungen", "Free+"],
        ["Menschliche Bewertung", "Mit Verdikt nicht einverstanden? Eigene Bewertung abgeben", "Free+"],
        ["Chrome Extension", "Rechtsklick zum Prüfen auf jeder Webseite", "Pro+"],
        ["Bookmarklets", "Ein-Klick-Verifizierung aus jedem Browser", "Alle Pläne"],
        ["API-Zugang", "REST API mit SSE-Streaming", "Team"],
        ["Dark Mode", "Volle Unterstützung für Dark/Light/System", "Alle Pläne"],
        ["Mehrsprachig", "Deutsche & englische Oberfläche", "Alle Pläne"],
        ["DSGVO-Kontrollen", "Daten exportieren, Konto löschen", "Alle Pläne"],
    ]
    add_table(doc, ["Feature", "Beschreibung", "Verfügbarkeit"], features_table)
    doc.add_page_break()

    # ── 4 ──
    doc.add_heading("4. Ausführliche Feature-Beschreibung", level=1)

    doc.add_heading("4.1 Faktencheck — Die Kern-Engine", level=2)
    add_para(doc, "Der Faktencheck ist das Herzstück von KLAR. Wenn Sie einen Text einreichen, passiert Folgendes:")
    add_bullet(doc, "KLAR nutzt Google Gemini Flash KI, um Ihren Text zu lesen und jeden Satz zu identifizieren, der eine faktische Behauptung enthält.", bold_prefix="Schritt 1 — Behauptungen extrahieren: ")
    add_bullet(doc, "Für jede Behauptung durchsucht KLAR zuerst Wikipedia und Wikidata (diese decken etwa 40 % typischer Behauptungen ab). Für die restlichen Behauptungen wird Serper.dev verwendet. Jede Behauptung erhält 3–5 Quellen.", bold_prefix="Schritt 2 — Belege suchen: ")
    add_bullet(doc, "Ein zweiter KI-Durchlauf bewertet jede Behauptung gegen die gesammelten Belege:", bold_prefix="Schritt 3 — Urteil: ")
    add_colored_bullet(doc, "16A34A", "Bestätigt (Grün) — Die Belege bestätigen die Richtigkeit dieser Behauptung")
    add_colored_bullet(doc, "DC2626", "Widerlegt (Rot) — Die Belege zeigen, dass diese Behauptung falsch ist")
    add_colored_bullet(doc, "64748B", "Unbestätigt (Grau) — Keine zuverlässige Quelle konnte gefunden werden. Das bedeutet nicht, dass die Behauptung falsch ist.")
    add_para(doc, "Jedes Urteil kommt mit einer verständlichen Erklärung und anklickbaren Quellenlinks.")

    doc.add_heading("4.2 AI-Erkennung", level=2)
    add_para(doc, "Diese Engine analysiert Schreibmuster, um festzustellen, ob ein Text von einem Menschen oder einer KI geschrieben wurde. Sie verwendet statistische Mathematik (NLP) — keine KI — was sie sofort und kostenlos macht.")
    add_para(doc, "Was gemessen wird:")
    add_bullet(doc, "KI tendiert zu sehr gleichförmigen Sätzen", bold_prefix="Satzvielfalt — ")
    add_bullet(doc, "KI verwendet ein engeres Wortspektrum", bold_prefix="Wortvielfalt — ")
    add_bullet(doc, "KI-Texte folgen vorhersagbareren Mustern", bold_prefix="Vorhersagbarkeit — ")
    add_bullet(doc, "KI verwendet oft bestimmte Phrasen und Strukturen", bold_prefix="Strukturmuster — ")
    add_para(doc, "Das Ergebnis ist ein Prozentscore (0–100 %) und ein Urteil: Menschlich, Wahrscheinlich menschlich, Gemischt, Wahrscheinlich KI oder KI-generiert.")

    doc.add_heading("4.3 Bias-Erkennung", level=2)
    add_para(doc, "Der Bias-Detektor liest Ihren gesamten Text und sucht nach Anzeichen von Voreingenommenheit.")
    add_para(doc, "Was erkannt wird:")
    add_bullet(doc, 'W\u00f6rter, die Emotionen ausl\u00f6sen sollen ("verheerend", "wundersam", "schockierend")', bold_prefix="Geladene Sprache \u2014 ")
    add_bullet(doc, "Versuche, den Leser fühlen statt denken zu lassen", bold_prefix="Emotionale Appelle — ")
    add_bullet(doc, "Darstellung nur einer Perspektive zu einem Thema", bold_prefix="Einseitige Darstellung — ")
    add_bullet(doc, "Ob der Text nach links, rechts oder neutral tendiert", bold_prefix="Politische Tendenz — ")
    add_para(doc, "Ergebnis: Bias-Score von 0 bis 100, Einstufung (Minimal, Gering, Mäßig, Erheblich, Extrem) und eine Zusammenfassung.")

    doc.add_heading("4.4 Plagiatscheck", level=2)
    add_para(doc, "Vergleicht Ihren Text gegen Webquellen. Ergebnis: Originalitäts-Prozentsatz, Liste übereinstimmender Quellen und Urteil (Original, Überwiegend original, Einige Überschneidungen, Erhebliche Überschneidungen, Wahrscheinlich plagiiert).")

    doc.add_heading("4.5 Qualitätsbewertung", level=2)
    add_para(doc, "Bewertet die Textqualität nach vier professionellen Frameworks:")
    add_bullet(doc, "Ist der Inhalt lückenlos und überschneidungsfrei organisiert?", bold_prefix="MECE — ")
    add_bullet(doc, "Könnte jemand die Argumente leicht widerlegen?", bold_prefix="Red Team — ")
    add_bullet(doc, "Wird der Hauptpunkt klar am Anfang genannt?", bold_prefix="BLUF — ")
    add_bullet(doc, "Was könnte schiefgehen, wenn jemand nach dieser Information handelt?", bold_prefix="Pre-Mortem — ")
    add_para(doc, "Ergebnis: Gesamtnote A bis F mit Einzelbewertungen und Verbesserungsvorschlägen.")

    doc.add_heading("4.6 Datei-Upload", level=2)
    add_para(doc, "Statt Text einzufügen können Sie direkt eine Datei hochladen: PDF, DOCX oder TXT. KLAR extrahiert den Text automatisch. Maximale Dateigröße: 10 MB (2 MB im Free-Plan).")

    doc.add_heading("4.7 URL-Verifizierung", level=2)
    add_para(doc, "Geben Sie die Webadresse eines Artikels ein. KLAR besucht die Seite, extrahiert den Hauptinhalt und analysiert ihn.")

    doc.add_heading("4.8 Trust Report", level=2)
    add_para(doc, "Nach der Analyse generiert KLAR einen umfassenden Trust Report:")
    add_bullet(doc, "Trust Score (0–100) basierend auf dem Verhältnis von bestätigten zu widerlegten Behauptungen (unbestätigte Behauptungen werden vom Score ausgeschlossen)")
    add_bullet(doc, "Farbcodierte Behauptungen mit Urteilen und Quellen")
    add_bullet(doc, "Hervorgehobene Textansicht — Ihr Originaltext mit markierten Behauptungen")
    add_bullet(doc, "Quellenlinks für jedes Urteil")
    add_bullet(doc, "PDF-Export des vollständigen Berichts")

    doc.add_heading("4.9 Dashboard & Prüfhistorie", level=2)
    add_para(doc, "Das Dashboard zeigt: Anzahl der Prüfungen, Trust-Score-Trends, aktuelle widerlegte Behauptungen und Schnellzugriff auf letzte Berichte. Die Prüfhistorie ist eine durchsuchbare, sortierbare Liste aller bisherigen Prüfungen.\n\nHinweis: Der Trust Score misst bestätigte vs. widerlegte Behauptungen. Unbestätigte Behauptungen (ohne Quellen für oder gegen) werden vom Score ausgeschlossen.")

    doc.add_heading("4.10 Menschliches Bewertungssystem", level=2)
    add_para(doc, "Wenn Sie mit KLARs Urteil nicht einverstanden sind, können Sie eine eigene Bewertung mit Kommentar abgeben.")
    doc.add_page_break()

    # ── 5 ──
    doc.add_heading("5. Wie die Technologie funktioniert", level=1)
    add_para(doc, "KLAR verwendet nicht eine einzelne KI, die alles macht. Stattdessen nutzt es eine Pipeline — eine Reihe spezialisierter Schritte, die jeweils eine Sache gut machen.")
    doc.add_heading("Die 10-Schritt-Pipeline", level=2)
    pipeline_steps = [
        ("1. AI-Erkennung", "Statistische Analyse bestimmt, ob der Text menschlich oder KI-generiert ist."),
        ("2. Qualitätsbewertung", "Text wird nach MECE, Red Team, BLUF und Pre-Mortem bewertet."),
        ("3. Behauptungen extrahieren", "Gemini Flash KI findet jede faktische Aussage als strukturierte Daten."),
        ("4. Belege suchen", "Wikipedia/Wikidata zuerst, dann Google. 3–5 Quellen pro Behauptung."),
        ("5. Claim-Qualitätsprüfung", "Wie spezifisch und prüfbar ist jede Behauptung?"),
        ("6. Halluzinations-Erkennung", "Typische KI-Erfindungsmuster werden identifiziert."),
        ("7. KI-Urteil", "Zweiter Gemini-Durchlauf bewertet jede Behauptung gegen die Belege."),
        ("8. Bias-Analyse", "NLP-basierte Erkennung von Tendenz und Einseitigkeit."),
        ("9. Plagiatscheck", "Text-Fingerprints werden gegen Web-Ergebnisse abgeglichen."),
        ("10. Kreuzreferenzierung", "Behauptungen werden quellenübergreifend auf Konsistenz geprüft."),
    ]
    for step_title, step_desc in pipeline_steps:
        add_bullet(doc, step_desc, bold_prefix=step_title + " — ")

    doc.add_heading("Welche KI verwendet KLAR?", level=2)
    add_para(doc, "KLAR nutzt Google Gemini 2.5 Flash in zwei Modi:")
    add_bullet(doc, "Zum Extrahieren von Behauptungen und Verstehen von Text", bold_prefix="Standardmodus — ")
    add_bullet(doc, "Gemini durchsucht das Internet in Echtzeit (\"Google Search Grounding\")", bold_prefix="Grounded-Search-Modus — ")
    add_para(doc, "KLAR verwendet KEINE Vektordatenbanken, Agent-Frameworks, RAG oder feinabgestimmte Modelle.")

    doc.add_heading("Die Analyse-Engines", level=2)
    engines_table = [
        ["Faktencheck", "KI + Suche", "Gemini extrahiert → Quellensuche → Gemini bewertet"],
        ["AI-Detektor", "Statist. Mathematik", "Satzmuster, Wortvielfalt, Vorhersagbarkeit"],
        ["Bias-Detektor", "Statist. Mathematik", "Emotionale Wörter, Framing, Quellenbalance"],
        ["Plagiatsprüfer", "Statist. Mathematik", "Text-Fingerprints vs. Web-Ergebnisse"],
        ["Qualitätsbewerter", "Statist. Mathematik", "MECE/Red Team/BLUF/Pre-Mortem"],
        ["Halluzinations-Detektor", "Statist. Mathematik", "KI-Erfindungsmuster"],
        ["Kreuzreferenz", "Logik", "Konsistenzprüfung über mehrere Quellen"],
    ]
    add_table(doc, ["Engine", "Typ", "Funktionsweise"], engines_table)
    doc.add_page_break()

    # ── 6 ──
    doc.add_heading("6. Die Web-Anwendung (UI-Leitfaden)", level=1)
    add_para(doc, "KLAR ist erreichbar unter klar-app.vercel.app. Die Oberfläche unterstützt Deutsch und Englisch mit einem Sprachumschalter im Header. Dark Mode, Light Mode und System-Erkennung sind eingebaut.")
    doc.add_heading("Hauptseiten", level=2)
    add_bullet(doc, "Moderne Startseite mit Erklärung zu KLAR und Schnellstart.", bold_prefix="Startseite — ")
    add_bullet(doc, 'Die Hauptanalyse-Oberfl\u00e4che. Text einf\u00fcgen, Datei hochladen oder URL eingeben. Analysemodi ausw\u00e4hlen. Auf "Analysieren" klicken.', bold_prefix='Pr\u00fcfen-Seite \u2014 ')
    add_bullet(doc, "Persönliche Übersicht mit Statistiken, Trust-Score-Trends und aktuellen Berichten.", bold_prefix="Dashboard — ")
    add_bullet(doc, "Durchsuchbare, sortierbare Liste aller bisherigen Prüfungen.", bold_prefix="Historie — ")
    add_bullet(doc, "Profilverwaltung, Datenexport (DSGVO) und Kontolöschung.", bold_prefix="Einstellungen — ")
    doc.add_heading("Der Prüfablauf", level=2)
    add_bullet(doc, "Navigieren Sie zur Prüfen-Seite")
    add_bullet(doc, "Fügen Sie Text ein (bis 10.000 Zeichen), laden Sie eine Datei hoch oder geben Sie eine URL ein")
    add_bullet(doc, "Wählen Sie Analysemodi (Faktencheck ist Standard)")
    add_bullet(doc, 'Klicken Sie auf "Analysieren"')
    add_bullet(doc, "Beobachten Sie die Ergebnisse in Echtzeit — jede Behauptung erscheint, sobald sie geprüft ist")
    add_bullet(doc, "Überprüfen Sie den Trust Report")
    add_bullet(doc, "Klicken Sie auf jede Behauptung für Details und Quellen")
    add_bullet(doc, "Bei Bedarf als PDF exportieren")
    doc.add_page_break()

    # ── 7 ──
    doc.add_heading("7. Die Chrome Extension", level=1)
    add_para(doc, "KLAR enthält eine Manifest-V3-Chrome-Extension, mit der Sie Text direkt auf jeder Webseite prüfen können.")
    doc.add_heading("Installation", level=2)
    add_bullet(doc, "API-Schlüssel auf der Einstellungsseite generieren")
    add_bullet(doc, "chrome://extensions öffnen")
    add_bullet(doc, '"Entwicklermodus" aktivieren')
    add_bullet(doc, '"Entpackte Erweiterung laden" klicken und den Ordner extension/ ausw\u00e4hlen')
    add_bullet(doc, "Das KLAR-Symbol in der Toolbar anklicken und den API-Schlüssel eingeben")
    doc.add_heading("Verwendung", level=2)
    add_bullet(doc, 'Text markieren \u2192 Rechtsklick \u2192 "Mit KLAR pr\u00fcfen"', bold_prefix='Rechtsklick-Pr\u00fcfung: ')
    add_bullet(doc, "Rechtsklick ohne Textauswahl prüft die gesamte Seite", bold_prefix="Ganzseitenprüfung: ")
    add_bullet(doc, "Ein schwebendes Panel zeigt Trust Score und Behauptungsurteile. Es werden nur Quellen angezeigt, die die KI tatsächlich referenziert hat.", bold_prefix="Inline-Ergebnisse: ")
    add_para(doc, "Kompatibel mit Chrome, Edge und allen Chromium-basierten Browsern.")
    doc.add_page_break()

    # ── 8 ──
    doc.add_heading("8. Browser-Bookmarklets", level=1)
    add_para(doc, "Für Browser ohne Extension-Unterstützung bietet KLAR Bookmarklets — kleine Buttons in der Lesezeichenleiste:")
    add_bullet(doc, "Sendet den markierten Text oder Seiteninhalt an KLAR", bold_prefix='"KLAR Verify" \u2014 ')
    add_bullet(doc, "Sendet die aktuelle Seiten-URL an KLAR", bold_prefix='"KLAR URL" \u2014 ')
    add_para(doc, "Keine Installation nötig. Funktioniert in Chrome, Firefox, Edge und Safari.")
    doc.add_page_break()

    # ── 9 ──
    doc.add_heading("9. API-Zugang für Entwickler", level=1)
    add_para(doc, "Entwickler können KLAR per REST API in eigene Anwendungen integrieren (Team-Plan).")
    add_bullet(doc, "POST-Anfrage an /api/verify mit Text und API-Schlüssel senden")
    add_bullet(doc, "Ergebnisse werden in Echtzeit per Server-Sent Events (SSE) gestreamt")
    add_bullet(doc, "Die Chrome Extension nutzt /api/extension/scan (JSON-Antwort)")
    add_para(doc, "Code-Beispiele in cURL, JavaScript und Python verfügbar.")
    doc.add_page_break()

    # ── 10 ──
    doc.add_heading("10. Preise & Nutzungslimits", level=1)
    pricing_table = [
        ["Guest", "Gratis (ohne Anmeldung)", "3/Monat", "2.000 Zeichen", "Nein", "Nur Faktencheck"],
        ["Free", "0 €/Monat", "10/Monat", "5.000 Zeichen", "Bis 2 MB", "Faktencheck, Bias, AI"],
        ["Pro", "9 €/Monat", "200/Monat", "10.000 Zeichen", "Bis 10 MB", "Alle 5 + Comprehensive"],
        ["Team", "29 €/Nutzer/Monat", "Unbegrenzt", "10.000 Zeichen", "Bis 10 MB", "Alle + API + Parallel"],
    ]
    add_table(doc, ["Plan", "Preis", "Monatl. Prüfungen", "Max. Text", "Datei-Upload", "Analysemodi"], pricing_table)
    doc.add_heading("Nutzungshinweise", level=2)
    add_bullet(doc, "Comprehensive-Analyse (alle 5 Engines gleichzeitig) zählt als 2–3 Prüfungen")
    add_bullet(doc, "Free- und Guest-Nutzer können jeweils eine Analyse gleichzeitig ausführen")
    add_bullet(doc, "Burst-Schutz: maximal 1–2 Prüfungen pro Minute in kostenlosen Plänen")
    add_bullet(doc, "Missbrauchsschutz: wiederholte Limitüberschreitung löst eskalierende Wartezeiten aus")
    doc.add_page_break()

    # ── 11 ──
    doc.add_heading("11. Datenschutz, Sicherheit & DSGVO", level=1)
    add_para(doc, "KLAR ist mit europäischen Datenschutzstandards im Kern gebaut:")
    add_bullet(doc, "Gebaut für die EU-Datenschutzgrundverordnung", bold_prefix="DSGVO-konform — ")
    add_bullet(doc, "Alle Daten in Frankfurt, Deutschland (Supabase EU)", bold_prefix="EU-Datenhaltung — ")
    add_bullet(doc, "Transparente KI-Entscheidungsfindung", bold_prefix="EU AI Act-bereit — ")
    add_bullet(doc, "Keine Tracking-Cookies, kein Datenverkauf, keine Drittanbieter-Analytics", bold_prefix="Keine Überwachung — ")
    add_bullet(doc, "Jeder Nutzer sieht nur seine eigenen Daten (Row Level Security)", bold_prefix="Datenisolation — ")
    add_bullet(doc, "Alles herunterladen, was KLAR über Sie gespeichert hat", bold_prefix="Datenexport — ")
    add_bullet(doc, "Konto und alle zugehörigen Daten dauerhaft löschen", bold_prefix="Kontolöschung — ")
    doc.add_heading("Sicherheitsmaßnahmen", level=2)
    add_bullet(doc, "Row Level Security (RLS) auf allen Supabase-Tabellen")
    add_bullet(doc, "API-Schlüssel nur serverseitig — nie im Browser sichtbar")
    add_bullet(doc, "Eingabevalidierung und -bereinigung aller Texteingaben")
    add_bullet(doc, "Prompt-Injection-Prävention mit isolierten KI-Aufrufen")
    add_bullet(doc, "CSRF-Schutz und HTTPS-Erzwingung")
    add_bullet(doc, "Rate Limiting pro Nutzer und IP-Adresse")
    doc.add_page_break()

    # ── 12 ──
    doc.add_heading("12. Für wen ist KLAR?", level=1)
    add_bullet(doc, "Behauptungen in Pressemitteilungen und Artikeln vor der Veröffentlichung prüfen", bold_prefix="Journalisten — ")
    add_bullet(doc, "KI-generierte Lernmaterialien und Arbeiten auf Richtigkeit prüfen", bold_prefix="Studierende — ")
    add_bullet(doc, "Daten validieren, Plagiate erkennen und akademische Integrität sichern", bold_prefix="Forscher — ")
    add_bullet(doc, "Blog-Beiträge und Social-Media-Inhalte auf Fakten prüfen", bold_prefix="Content Creator — ")
    add_bullet(doc, "KI-generierte Berichte, Angebote und Kundendokumentation prüfen", bold_prefix="Unternehmen — ")
    add_bullet(doc, "Verträge, Compliance-Dokumente und juristische Recherchen verifizieren", bold_prefix="Rechts- & Beratungsfirmen — ")
    add_bullet(doc, "Lehrmaterialien und Studierendenarbeiten auf Richtigkeit und Originalität prüfen", bold_prefix="Bildungseinrichtungen — ")
    add_bullet(doc, "Jeder, der wissen will, ob das, was er liest, tatsächlich stimmt", bold_prefix="Alle — ")
    doc.add_page_break()

    # ── 13 ──
    doc.add_heading("13. Zukünftige Entwicklung", level=1)
    add_para(doc, "KLAR wird aktiv weiterentwickelt. Geplant sind:")
    add_bullet(doc, "Echtzeit-Überwachung von Nachrichtenquellen auf Fehlinformationen", bold_prefix="Desinformations-Monitoring — ")
    add_bullet(doc, "Unterstützung größerer Teams mit Admin-Dashboards", bold_prefix="Team-Zusammenarbeit — ")
    add_bullet(doc, "Individuelle Verifizierungs-Workflows per API", bold_prefix="Erweiterte API-Features — ")
    add_bullet(doc, "Extensions für Firefox, Safari und andere Browser", bold_prefix="Weitere Browser-Extensions — ")
    add_bullet(doc, "Analyse von Bildern und Videos auf manipulierten Inhalt", bold_prefix="Visuelle Inhaltsverifizierung — ")
    add_bullet(doc, "Noch mehr Oberflächen-Sprachen über Deutsch und Englisch hinaus", bold_prefix="Weitere Sprachen — ")
    add_bullet(doc, "Integration mit Slack, Teams und anderen Kommunikationstools", bold_prefix="Workspace-Integrationen — ")
    doc.add_page_break()

    # ── 14 ──
    doc.add_heading("14. Schnellstart-Anleitung", level=1)
    doc.add_heading("Option A: Sofort-Prüfung (ohne Anmeldung)", level=2)
    add_bullet(doc, "Gehen Sie zu klar-app.vercel.app")
    add_bullet(doc, 'Klicken Sie auf "Pr\u00fcfen" in der Navigation')
    add_bullet(doc, "Fügen Sie Text ein")
    add_bullet(doc, 'Klicken Sie auf "Analysieren"')
    add_bullet(doc, "Überprüfen Sie Ihren Trust Report")
    add_para(doc, "Sie erhalten 3 kostenlose Prüfungen pro Monat als Gast.")
    doc.add_heading("Option B: Kostenloses Konto erstellen", level=2)
    add_bullet(doc, 'Klicken Sie auf "Registrieren" und geben Sie E-Mail + Passwort ein (oder Google nutzen)')
    add_bullet(doc, "Bestätigen Sie Ihre E-Mail-Adresse")
    add_bullet(doc, "Sie erhalten 10 Prüfungen/Monat, Dashboard und Prüfhistorie")
    doc.add_heading("Option C: Chrome Extension installieren", level=2)
    add_bullet(doc, "API-Schlüssel in den Einstellungen generieren (kostenloses Konto erforderlich)")
    add_bullet(doc, "Extension im Chrome-Entwicklermodus laden")
    add_bullet(doc, "Rechtsklick auf beliebigen Text auf jeder Webseite zum sofortigen Prüfen")
    doc.add_page_break()

    # ── 15 ──
    doc.add_heading("15. Häufige Fragen", level=1)
    faqs = [
        ("Wie genau ist KLAR?", 'Die Genauigkeit h\u00e4ngt von der Verf\u00fcgbarkeit der Quellen ab. Behauptungen, die gegen Wikipedia und zuverl\u00e4ssige Webquellen gepr\u00fcft werden k\u00f6nnen, erhalten hochgenaue Urteile. Behauptungen \u00fcber sehr aktuelle Ereignisse oder Nischenthemen k\u00f6nnen als "Unbest\u00e4tigt" eingestuft werden. Unbest\u00e4tigt bedeutet nicht falsch \u2014 es bedeutet nur, dass keine Quelle gefunden werden konnte.'),
        ("Speichert KLAR meinen Originaltext?", "KLAR speichert die extrahierten Behauptungen und ihre Urteile, damit Sie auf Ihre Historie zugreifen können. Alle Daten werden in Frankfurt (EU) mit Row Level Security gespeichert."),
        ("Kann KLAR Texte in anderen Sprachen prüfen?", "Die Analyse-Engine verarbeitet Texte in vielen Sprachen, da die Quellensuche sprachübergreifend funktioniert. Die Benutzeroberfläche ist in Deutsch und Englisch verfügbar."),
        ("Sind meine Daten sicher?", "Ja. KLAR ist DSGVO-konform, speichert alle Daten in Frankfurt (EU), verwendet Row Level Security und nutzt keine Tracking-Cookies oder verkauft Daten an Dritte."),
        ("Wie unterscheidet sich KLAR davon, ChatGPT zum Faktenprüfen zu bitten?", "ChatGPT kann das Internet nicht in Echtzeit durchsuchen und erfindet oft Quellen. KLAR durchsucht tatsächlich Live-Quellen (Wikipedia, Google) und liefert anklickbare Links zu jeder verwendeten Quelle."),
        ("Kann ich KLAR kommerziell nutzen?", "Ja. Die Pro- und Team-Pläne sind für professionelle und kommerzielle Nutzung konzipiert."),
    ]
    for q, a in faqs:
        doc.add_heading(q, level=3)
        add_para(doc, a)

    add_para(doc, "")
    p = doc.add_paragraph("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— Ende des Handbuchs —")
    r.font.name = FONT_NAME
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r.italic = True


# ─────────────────────────────────────────────────────────────────────────
# GENERATE
# ─────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # English
    doc_en = Document()
    for section in doc_en.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    build_english(doc_en)
    en_path = os.path.join(OUTPUT_DIR, "KLAR_User_Manual_EN.docx")
    doc_en.save(en_path)
    print(f"Created: {en_path}")

    # German
    doc_de = Document()
    for section in doc_de.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    build_german(doc_de)
    de_path = os.path.join(OUTPUT_DIR, "KLAR_Benutzerhandbuch_DE.docx")
    doc_de.save(de_path)
    print(f"Created: {de_path}")

    print("Done!")

